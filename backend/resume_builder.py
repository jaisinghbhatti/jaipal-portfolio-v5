"""
Resume Builder API endpoints with AI integration using Gemini 2.5 Flash
"""

from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
import os
import logging
import io
import re
import json
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/resume-builder", tags=["resume-builder"])

# Pydantic Models
class AnalyzeRequest(BaseModel):
    resumeText: str
    resumeParsed: Optional[Dict[str, Any]] = None
    jobDescription: str

class AnalyzeResponse(BaseModel):
    matchScore: int
    missingKeywords: List[str]

class OptimizeRequest(BaseModel):
    resumeText: str
    resumeParsed: Optional[Dict[str, Any]] = None
    jobDescription: str
    tone: str = "executive"

class OptimizeResponse(BaseModel):
    optimizedResume: str
    coverLetter: str
    newMatchScore: Optional[int] = None


# Helper functions
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file using multiple methods"""
    text = ""
    
    # Try PyPDF2 first
    try:
        from PyPDF2 import PdfReader
        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        text = text.strip()
        if text and len(text) > 20:
            logger.info(f"PyPDF2 extracted {len(text)} characters")
            return text
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed: {e}")
    
    # Fallback to pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        text = text.strip()
        if text:
            logger.info(f"pdfplumber extracted {len(text)} characters")
            return text
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")
    
    if not text:
        logger.error("Could not extract text from PDF with any method")
        raise HTTPException(status_code=400, detail="Could not extract text from PDF. The file may be image-based or corrupted. Please try pasting your resume text instead.")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        
        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        text = "\n".join(paragraphs)
        if not text:
            logger.warning("DOCX extraction returned empty text")
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}")


def parse_resume_text(text: str) -> Dict[str, Any]:
    """Basic parsing of resume text into structured data"""
    parsed = {
        "fullName": "",
        "contact": "",
        "summary": "",
        "experience": [],
        "education": [],
        "skills": []
    }
    
    # Try to extract name (usually first line or lines)
    lines = text.split('\n')
    for line in lines[:5]:
        line = line.strip()
        if line and len(line) < 50 and not any(char.isdigit() for char in line[:5]):
            # Likely a name
            if not parsed["fullName"]:
                parsed["fullName"] = line
                break
    
    # Extract email
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    emails = re.findall(email_pattern, text)
    if emails:
        parsed["contact"] = emails[0]
    
    # Extract skills section (basic)
    skills_keywords = ["skills", "technologies", "tools", "proficiencies"]
    for i, line in enumerate(lines):
        if any(kw in line.lower() for kw in skills_keywords):
            # Get next few lines as skills
            skill_lines = []
            for j in range(i+1, min(i+10, len(lines))):
                if lines[j].strip() and len(lines[j].strip()) < 100:
                    skill_lines.append(lines[j].strip())
                else:
                    break
            parsed["skills"] = skill_lines
            break
    
    return parsed


async def call_gemini_ai(prompt: str, system_message: str = "") -> str:
    """Call Gemini 2.5 Flash using emergentintegrations"""
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        api_key = os.environ.get('EMERGENT_LLM_KEY')
        if not api_key:
            raise HTTPException(status_code=500, detail="AI service not configured")
        
        chat = LlmChat(
            api_key=api_key,
            session_id=f"resume-builder-{os.urandom(8).hex()}",
            system_message=system_message or "You are an expert resume writer and career coach."
        ).with_model("gemini", "gemini-2.5-flash")
        
        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)
        
        return response
        
    except Exception as e:
        logger.error(f"Gemini AI error: {e}")
        raise HTTPException(status_code=500, detail=f"AI service error: {str(e)}")


# API Endpoints
@router.post("/parse")
async def parse_document(
    file: UploadFile = File(...),
    type: str = Form(...)
):
    """Parse uploaded PDF or DOCX file"""
    try:
        file_bytes = await file.read()
        filename = file.filename.lower() if file.filename else ""
        content_type = file.content_type or ""
        
        logger.info(f"Parse request - filename: {filename}, content_type: {content_type}, size: {len(file_bytes)}")
        
        # Check by filename extension or content type
        is_pdf = filename.endswith('.pdf') or 'pdf' in content_type
        is_docx = filename.endswith('.docx') or 'wordprocessingml' in content_type or 'msword' in content_type
        
        if is_pdf:
            text = extract_text_from_pdf(file_bytes)
        elif is_docx:
            text = extract_text_from_docx(file_bytes)
        else:
            logger.warning(f"Unsupported file type - filename: {filename}, content_type: {content_type}")
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF or DOCX.")
        
        if not text or len(text.strip()) < 10:
            logger.warning(f"Extracted text is too short: {len(text) if text else 0} chars")
            raise HTTPException(status_code=400, detail="Could not extract text from file. The file may be empty or image-based.")
        
        result = {"text": text}
        
        # If it's a resume, also parse it
        if type == "resume":
            result["parsed"] = parse_resume_text(text)
        
        logger.info(f"Successfully parsed file, extracted {len(text)} characters")
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Parse error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {str(e)}")


@router.post("/analyze", response_model=AnalyzeResponse)
async def analyze_resume(request: AnalyzeRequest):
    """Analyze resume against job description to calculate match score"""
    try:
        prompt = f"""You are a strict ATS keyword analyzer. Compare the resume against the job description.

STEP 1: Extract ALL important keywords, skills, tools, technologies, certifications, and phrases from the JOB DESCRIPTION.
STEP 2: Check which of those keywords are MISSING or NOT EXPLICITLY MENTIONED in the RESUME.
STEP 3: Calculate a match score (0-100) based on how many JD keywords appear in the resume.

CRITICAL RULES:
- You MUST find at least 3-5 missing keywords. No resume is a 100% match.
- Look for specific tools, technologies, soft skills, certifications, and industry terms in the JD.
- A keyword is "missing" if it does NOT appear in the resume text (even synonyms count as missing if the exact term is absent).
- Be strict: if the JD says "Python" and the resume says "programming", Python is still missing.
- Never return an empty missingKeywords list.

RESUME:
{request.resumeText}

JOB DESCRIPTION:
{request.jobDescription}

Respond in this exact JSON format ONLY, no other text or explanation:
{{
    "matchScore": <number 0-100>,
    "missingKeywords": ["keyword1", "keyword2", "keyword3", ...]
}}
"""
        
        system_message = "You are a strict ATS keyword matching engine. Always find gaps between resumes and job descriptions. Respond with valid JSON only, no markdown."
        
        response = await call_gemini_ai(prompt, system_message)
        
        # Parse the response
        # Clean the response - remove markdown code blocks if present
        clean_response = response.strip()
        if clean_response.startswith("```"):
            # Handle ```json ... ``` blocks
            parts = clean_response.split("```")
            for part in parts:
                part = part.strip()
                if part.startswith("json"):
                    part = part[4:].strip()
                if part.startswith("{"):
                    clean_response = part
                    break
        clean_response = clean_response.strip()
        
        # Remove any trailing text after the JSON object
        brace_count = 0
        end_idx = 0
        for i, ch in enumerate(clean_response):
            if ch == '{':
                brace_count += 1
            elif ch == '}':
                brace_count -= 1
                if brace_count == 0:
                    end_idx = i + 1
                    break
        if end_idx > 0:
            clean_response = clean_response[:end_idx]
        
        try:
            result = json.loads(clean_response)
        except json.JSONDecodeError:
            logger.warning(f"Failed to parse AI response as JSON: {response[:500]}")
            result = {
                "matchScore": 50,
                "missingKeywords": ["Unable to analyze - please try again"]
            }
        
        keywords = result.get("missingKeywords", [])
        if not keywords or (len(keywords) == 0):
            keywords = ["Review job description for specific requirements"]
        
        return AnalyzeResponse(
            matchScore=min(100, max(0, int(result.get("matchScore", 50)))),
            missingKeywords=keywords[:10]
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Analysis error: {e}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")


@router.post("/optimize", response_model=OptimizeResponse)
async def optimize_resume(request: OptimizeRequest):
    """Optimize resume with AI and generate cover letter"""
    try:
        # Define tone-specific instructions
        tone_instructions = {
            "executive": """Use high-level strategic language. 
            Preferred verbs: Spearheaded, Orchestrated, Leveraged, Directed, Championed
            Focus on leadership, strategy, and organizational impact.""",
            
            "disruptor": """Use punchy, action-oriented language.
            Preferred verbs: Built, Scaled, Accelerated, Disrupted, Transformed
            Focus on growth, innovation, and bold achievements.""",
            
            "human": """Use friendly, approachable language.
            Preferred verbs: Collaborated, Supported, Mentored, Facilitated, Nurtured
            Focus on teamwork, cultural fit, and people skills."""
        }
        
        tone_guide = tone_instructions.get(request.tone, tone_instructions["executive"])
        
        # Optimize Resume
        resume_prompt = f"""You are an expert resume writer. Rewrite this resume to be optimized for the given job description.

CRITICAL RULES:
1. For EVERY bullet point, include quantifiable metrics. If the original doesn't have numbers, add placeholders like [X%], [Y amount], [$Z], [N employees]
2. Apply the AIDA framework (Attention, Interest, Desire, Action) to the summary section
3. Incorporate missing keywords from the job description naturally
4. Keep the same structure but enhance the content

TONE STYLE:
{tone_guide}

ORIGINAL RESUME:
{request.resumeText}

TARGET JOB DESCRIPTION:
{request.jobDescription}

Provide the optimized resume as clean, formatted text. Use proper sections (SUMMARY, EXPERIENCE, EDUCATION, SKILLS, etc.)
Do NOT include any explanations or notes - just the optimized resume content."""

        optimized_resume = await call_gemini_ai(resume_prompt)
        
        # Generate Cover Letter
        cover_letter_prompt = f"""Generate a professional 3-paragraph cover letter based on this resume and job description.

STRUCTURE:
Paragraph 1 (HOOK): Open with attention-grabbing statement + job title being applied for
Paragraph 2 (CONNECTION): Connect 2-3 specific JD requirements with matching resume achievements
Paragraph 3 (CTA): Express enthusiasm + clear call to action

TONE STYLE:
{tone_guide}

RESUME:
{request.resumeText}

JOB DESCRIPTION:
{request.jobDescription}

Write ONLY the cover letter content, starting with "Dear Hiring Manager," and ending with a professional sign-off.
Do NOT include any explanations or notes - just the cover letter."""

        cover_letter = await call_gemini_ai(cover_letter_prompt)
        
        # Calculate new match score
        new_match_prompt = f"""Calculate a match score (0-100) for this optimized resume against the job description.
Consider keyword density, skill alignment, and experience relevance.

OPTIMIZED RESUME:
{optimized_resume}

JOB DESCRIPTION:
{request.jobDescription}

Respond with ONLY a single number between 0 and 100, nothing else."""

        try:
            new_score_response = await call_gemini_ai(new_match_prompt)
            new_match_score = int(re.search(r'\d+', new_score_response).group())
            new_match_score = min(100, max(0, new_match_score))
        except:
            new_match_score = None
        
        return OptimizeResponse(
            optimizedResume=optimized_resume.strip(),
            coverLetter=cover_letter.strip(),
            newMatchScore=new_match_score
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Optimization error: {e}")
        raise HTTPException(status_code=500, detail=f"Optimization failed: {str(e)}")
